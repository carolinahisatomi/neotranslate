import difflib
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

def adicionar_secao_formatada(doc, titulo, estilo_base):
    doc.add_page_break()
    par = doc.add_paragraph()
    run = par.add_run(titulo.upper())
    run.bold = True
    run.font.size = Pt(estilo_base['tamanho'] + 2)
    run.font.name = estilo_base['fonte']
    par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def titulo_ja_existe(doc, titulo_alvo, limiar=0.75):
    titulos_doc = [
        p.text.strip() for p in doc.paragraphs
        if p.style.name.startswith("Heading") or (p.text.isupper() and len(p.text.split()) <= 6)
    ]
    parecidos = difflib.get_close_matches(titulo_alvo.upper(), titulos_doc, n=1, cutoff=limiar)
    return bool(parecidos)

def adicionar_secoes_padrao(doc, estilo_base,
                            incluir_advertencias=True,
                            incluir_especificacoes=True,
                            incluir_orientacoes=True):
    if incluir_advertencias and not titulo_ja_existe(doc, "ADVERTÊNCIAS"):
        adicionar_secao_formatada(doc, "ADVERTÊNCIAS", estilo_base)

    if incluir_especificacoes and not titulo_ja_existe(doc, "ESPECIFICAÇÕES TÉCNICAS"):
        adicionar_secao_formatada(doc, "ESPECIFICAÇÕES TÉCNICAS", estilo_base)

    if incluir_orientacoes and not titulo_ja_existe(doc, "ORIENTAÇÕES"):
        adicionar_secao_formatada(doc, "ORIENTAÇÕES", estilo_base)
