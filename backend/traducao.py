import os
import csv
import time
import hashlib
import difflib
from io import BytesIO
from docx import Document
from docx.shared import Pt
import streamlit as st

from backend.memoria import carregar_memoria, salvar_memoria
from backend.tradutor import traduzir_com_memoria
from backend.inmetro import adicionar_textos_inmetro, adicionar_texto_reciclagem
from formatador_docx import adicionar_secoes_padrao

TIPOS_INVERSOR = {
    "Inversor Off-grid",
    "Inversor On-grid",
    "Inversor Carregador",
    "Inversor Híbrido"
}

def extrair_modelo(texto):
    import re
    padroes = [
        r"Model(?: No\.| Number)?[:\s]+([A-Za-z0-9\-_/]+)",
        r"Modelo[:\s]+([A-Za-z0-9\-_/]+)"
    ]
    for padrao in padroes:
        match = re.search(padrao, texto, re.IGNORECASE)
        if match:
            return match.group(1).replace(" ", "_")
    return "modelo"

def traduzir_tabelas(doc_origem, doc_final, memoria, estilo_base, aplicar_glossario_func=None,
                     progresso=None, itens_processados=0, total_itens=1):
    for tabela in doc_origem.tables:
        nova_tabela = doc_final.add_table(rows=0, cols=len(tabela.columns))
        for linha in tabela.rows:
            nova_linha = nova_tabela.add_row().cells
            for j, celula in enumerate(linha.cells):
                texto = celula.text.strip()
                if not texto:
                    nova_linha[j].text = ""
                    itens_processados += 1
                    if progresso:
                        progresso.progress(itens_processados / total_itens, text="Traduzindo conteúdo...")
                    continue

                if aplicar_glossario_func:
                    texto = aplicar_glossario_func(texto)

                traducao, _ = traduzir_com_memoria(texto, memoria)
                nova_linha[j].text = traducao

                for par in nova_linha[j].paragraphs:
                    for run in par.runs:
                        run.font.name = estilo_base['fonte']
                        run.font.size = Pt(estilo_base['tamanho'])

                itens_processados += 1
                if progresso:
                    progresso.progress(itens_processados / total_itens, text="Traduzindo conteúdo...")
                time.sleep(0.5)

def traduzir_docx_com_tudo(arquivo, exige_inmetro, tipo_equipamento, aplicar_glossario_func):
    doc_origem = Document(arquivo)
    doc_final = Document()
    memoria = carregar_memoria()

    estilo_base = {"fonte": "Arial", "tamanho": 10}
    for par in doc_origem.paragraphs:
        if par.text.strip() and par.runs:
            estilo_base["fonte"] = par.runs[0].font.name or "Arial"
            estilo_base["tamanho"] = par.runs[0].font.size.pt if par.runs[0].font.size else 10
            break

    total_paragrafos = len([p for p in doc_origem.paragraphs if p.text.strip()])
    total_celulas = sum(len(linha.cells) for tabela in doc_origem.tables for linha in tabela.rows)
    total_itens = total_paragrafos + total_celulas
    progresso = st.progress(0, text="Iniciando tradução...")
    itens_processados = 0

    modelo = extrair_modelo(" ".join(p.text for p in doc_origem.paragraphs))
    nome_docx = f"manual_{modelo}_traduzido.docx"
    nome_csv = f"memoria_traducao/memoria_traducao_{modelo}.csv"

    aprovar_tudo = st.checkbox("✅ Aprovar todas as traduções automaticamente", value=False)

    traducoes_corrigidas = []
    for i, par in enumerate(doc_origem.paragraphs):
        texto = par.text.strip()
        if not texto:
            continue

        traducao, _ = traduzir_com_memoria(texto, memoria)
        hash_texto = hashlib.md5(texto.encode()).hexdigest()
        chave = f"par_{i}_{hash_texto}"

        if aprovar_tudo:
            corrigido = traducao
            st.session_state[chave] = corrigido  # garante que fique disponível se checkbox foi marcado
        else:
            st.markdown(f"**Original:** {texto}")
            corrigido = st.text_area(
                "Tradução corrigida",
                key=chave,
                value=traducao if chave not in st.session_state else st.session_state[chave]
            )

        p = doc_final.add_paragraph(st.session_state[chave])
        run = p.runs[0]
        run.font.name = estilo_base["fonte"]
        run.font.size = Pt(estilo_base["tamanho"])
        traducoes_corrigidas.append((texto, traducao, st.session_state[chave]))
        itens_processados += 1
        progresso.progress(itens_processados / total_itens, text="Traduzindo conteúdo...")

    if exige_inmetro and tipo_equipamento in TIPOS_INVERSOR:
        adicionar_textos_inmetro(doc_final, estilo_base)

    adicionar_texto_reciclagem(doc_final, tipo_equipamento, estilo_base)
    if exige_inmetro:
        adicionar_secoes_padrao(doc_final, estilo_base)

    traduzir_tabelas(
        doc_origem=doc_origem,
        doc_final=doc_final,
        memoria=memoria,
        estilo_base=estilo_base,
        aplicar_glossario_func=aplicar_glossario_func,
        progresso=progresso,
        itens_processados=itens_processados,
        total_itens=total_itens
    )

    salvar_memoria(os.path.basename(nome_csv), traducoes_corrigidas)

    buffer_docx = BytesIO()
    doc_final.save(buffer_docx)
    buffer_docx.seek(0)

    st.download_button(
        label="📥 Baixar Tradução Final (.docx)",
        data=buffer_docx,
        file_name=nome_docx,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    st.success("Arquivo CSV salvo na pasta 'memoria_traducao'")
    progresso.empty()

# Formatação TOTAL

TITULOS_INMETRO = {
    "Especificações Técnicas",
    "Requisitos de Instalação",
    "Instruções de Segurança",
    "Informações do Fabricante",
    "Garantia",
    "Reciclagem e Descarte"
}

def encontrar_titulo_semelhante(texto, titulos_alvo, limiar=0.75):
    correspondencias = difflib.get_close_matches(texto, titulos_alvo, n=1, cutoff=limiar)
    return correspondencias[0] if correspondencias else None

def traduzir_docx_preservando_layout(arquivo, exige_inmetro, tipo_equipamento, aplicar_glossario_func):
    doc = Document(arquivo)
    memoria = carregar_memoria()

    estilo_base = {"fonte": "Arial", "tamanho": 10}
    for par in doc.paragraphs:
        if par.text.strip() and par.runs:
            estilo_base["fonte"] = par.runs[0].font.name or "Arial"
            estilo_base["tamanho"] = par.runs[0].font.size.pt if par.runs[0].font.size else 10
            break

    modelo = "modelo"
    for par in doc.paragraphs:
        if "model" in par.text.lower() or "modelo" in par.text.lower():
            modelo = par.text.split()[-1].replace(":", "").strip()
            break

    nome_docx = f"manual_{modelo}_traduzido.docx"
    nome_csv = f"memoria_traducao/memoria_traducao_{modelo}.csv"
    progresso = st.progress(0, text="Traduzindo manual...")

    aprovar_tudo = st.checkbox("✅ Aprovar todas as traduções automaticamente", value=False)

    traducoes_corrigidas = []
    total_paragrafos = len([p for p in doc.paragraphs if p.text.strip()])
    total_celulas = sum(len(linha.cells) for tabela in doc.tables for linha in tabela.rows)
    total_itens = total_paragrafos + total_celulas
    itens_processados = 0

    # Tradução de parágrafos
    for i, par in enumerate(doc.paragraphs):
        texto = par.text.strip()
        if not texto:
            continue

        traducao, _ = traduzir_com_memoria(texto, memoria)
        hash_texto = hashlib.md5(texto.encode()).hexdigest()
        chave = f"par_{i}_{hash_texto}"

        if aprovar_tudo:
            st.session_state[chave] = traducao
        else:
            st.markdown(f"**Original:** {texto}")
            if chave not in st.session_state:
                st.session_state[chave] = traducao
            corrigido = st.text_area(
                "Tradução corrigida",
                value=st.session_state[chave],
                key=chave,
            )

        par.text = st.session_state[chave]
        traducoes_corrigidas.append((texto, traducao, st.session_state[chave]))
        itens_processados += 1
        progresso.progress(itens_processados / total_itens, text="Traduzindo parágrafos...")

    # Tradução de tabelas
    for t_idx, tabela in enumerate(doc.tables):
        for r_idx, linha in enumerate(tabela.rows):
            for c_idx, celula in enumerate(linha.cells):
                for p_idx, par in enumerate(celula.paragraphs):
                    texto = par.text.strip()
                    if not texto:
                        continue
                    traducao, _ = traduzir_com_memoria(texto, memoria)
                    hash_texto = hashlib.md5(texto.encode()).hexdigest()
                    chave = f"celula_{t_idx}_{r_idx}_{c_idx}_{p_idx}_{hash_texto}"
                    if aprovar_tudo:
                        st.session_state[chave] = traducao
                    else:
                        st.markdown(f"**Original (tabela):** {texto}")
                        if chave not in st.session_state:
                            st.session_state[chave] = traducao
                        corrigido = st.text_area(
                            "Tradução corrigida (tabela)",
                            value=st.session_state[chave],
                            key=chave,
                        )
                    par.text = st.session_state[chave]
                    traducoes_corrigidas.append((texto, traducao, st.session_state[chave]))
                    itens_processados += 1
                    progresso.progress(itens_processados / total_itens, text="Traduzindo tabelas...")

    # Inserção dos textos do INMETRO (com verificação de títulos semelhantes)
    if exige_inmetro and tipo_equipamento in TIPOS_INVERSOR:
        titulos_existentes = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading")]
        inseridos = set()

        for titulo_necessario in TITULOS_INMETRO:
            similar = encontrar_titulo_semelhante(titulo_necessario, titulos_existentes)
            if similar:
                # Insere abaixo da seção similar
                for idx, par in enumerate(doc.paragraphs):
                    if par.text.strip() == similar:
                        doc.paragraphs.insert(idx + 1, Document().add_paragraph(f"[Texto obrigatório: {titulo_necessario}]"))
                        inseridos.add(titulo_necessario)
                        break
            else:
                # Insere normalmente no final se não encontrar seção parecida
                doc.add_paragraph(titulo_necessario, style="Heading 1")
                doc.add_paragraph(f"[Texto obrigatório: {titulo_necessario}]")
                inseridos.add(titulo_necessario)

    adicionar_texto_reciclagem(doc, tipo_equipamento, estilo_base)
    if exige_inmetro:
        adicionar_secoes_padrao(doc, estilo_base)

    salvar_memoria(os.path.basename(nome_csv), traducoes_corrigidas)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="📥 Baixar Tradução Final (.docx)",
        data=buffer,
        file_name=nome_docx,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    st.success("Tradução concluída e memória salva.")
    progresso.empty()