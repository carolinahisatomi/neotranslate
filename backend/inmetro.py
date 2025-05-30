import json
from docx.shared import Pt
from docx import Document

# Espera-se que o JSON já esteja carregado no app principal
TEXTOS_DESCARTE = {}

def configurar_textos_descarte(json_dict):
    global TEXTOS_DESCARTE
    TEXTOS_DESCARTE = json_dict

def adicionar_textos_inmetro(doc: Document, estilo_base):
    def par(texto, negrito=False):
        p = doc.add_paragraph()
        run = p.add_run(texto)
        run.bold = negrito
        run.font.name = estilo_base['fonte']
        run.font.size = Pt(estilo_base['tamanho'])

    par("IMPORTANTE LER COM ATENÇÃO E GUARDAR PARA EVENTUAIS CONSULTAS", negrito=True)
    par("Distribuidor Oficial no Brasil\nAssistência Técnica no Brasil\nNEOSOLAR ENERGIA LTDA")
    par("CNPJ 12.420.339/0001-26 | Rua Morgado de Mateus, 516 – São Paulo – SP")
    par("www.neosolar.com.br | contato@neosolar.com.br")
    par("Fone SAC (11) 4328-5113 | WhatsApp (11) 99935-4535")
    par("Atenção: A instalação desse equipamento deve obedecer às normas técnicas vigentes para instalação elétrica fotovoltaica (NBR 16690) e gestão de riscos de incêndios em sistemas fotovoltaicos (IEC 63226).")
    par("Atenção: necessita de dispositivo externo de proteção")
    par("Atenção: necessita de dispositivo de interrupção multipolar para desconexão dos condutores de corrente")
    par("Atenção: necessita de dispositivo de corrente residual (DR) externo, adequado para proteção contra choque elétrico, de acordo com a norma ABNT NBR 5410")
    par("Atenção: é expressamente recomendada a utilização de métodos, sistemas ou dispositivos de desligamento rápido no circuito c.c. que garantam a segurança em situações de combate à incêndio.")
    par("Compatibilidade com sistemas de desligamento rápido")

def adicionar_texto_reciclagem(doc: Document, tipo: str, estilo_base):
    texto = TEXTOS_DESCARTE.get(tipo, "")
    if texto:
        p = doc.add_paragraph()
        run = p.add_run(f"\n{texto}")
        run.font.name = estilo_base['fonte']
        run.font.size = Pt(estilo_base['tamanho'])